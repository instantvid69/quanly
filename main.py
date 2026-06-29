import flet as ft
import pandas as pd
import unicodedata
import re
import os
from datetime import datetime, date
from dateutil.relativedelta import relativedelta

def normalize_text(text: str) -> str:
    if not text:
        return ""
    s = str(text).replace("Đ", "D").replace("đ", "d")
    return "".join(
        c for c in unicodedata.normalize("NFD", s)
        if unicodedata.category(c) != "Mn"
    ).lower().strip()

def clean_unit_name(unit_str: str) -> str:
    if not unit_str or str(unit_str).strip().lower() in ['nan', '', '-', 'nat']:
        return "Chưa cập nhật"
    return str(unit_str).strip()

def clean_cell_to_date_str(val) -> str:
    if pd.isna(val) or str(val).strip().lower() in ['nan', '', 'nat', '-', '.', '0']:
        return ""
    if isinstance(val, (datetime, date, pd.Timestamp)):
        return val.strftime('%d/%m/%Y')

    raw_s = str(val).strip().split()[0]
    s_slashes = raw_s.replace('-', '/').replace('.', '/')

    for fmt in ('%d/%m/%Y', '%Y/%m/%d', '%m/%d/%Y', '%d/%m/%y', '%Y%m%d'):
        try:
            test_str = raw_s if fmt == '%Y%m%d' else s_slashes
            dt = datetime.strptime(test_str, fmt)
            return dt.strftime('%d/%m/%Y')
        except ValueError:
            continue
    return str(val).strip()

def find_column_by_keywords(columns, keywords) -> str | None:
    for kw in keywords:
        for col in columns:
            if kw == normalize_text(col):
                return col
    for kw in keywords:
        for col in columns:
            if kw in normalize_text(col):
                return col
    return None

class AllowanceDashboardApp:
    def __init__(self, page: ft.Page):
        self.page = page
        self.page.title = "Hệ thống Quản lý Thu hút & Phụ cấp"
        self.page.theme_mode = ft.ThemeMode.LIGHT
        self.page.bgcolor = ft.Colors.BLUE_GREY_50
        self.page.padding = 20

        # Dữ liệu gốc và dữ liệu sau bộ lọc
        self.master_data = {}
        self.filtered_data = {}
        self.extra_headers = []

        # Các biến phục vụ phân trang
        self.per_page = 50          
        self.page_number = 1        
        self.stats_page_number = 1  
        self.adjust_page_number = 1
        self.stats_filtered_items = [] 

        # UI cho Tab 1: Danh Sách
        self.search_field = ft.TextField(
            label="Tìm kiếm theo tên, năm sinh, đơn vị, số hiệu...",
            prefix_icon=ft.Icons.SEARCH,
            on_submit=self.filter_allowance_list,
            border_radius=12,
            bgcolor=ft.Colors.WHITE,
            border_color=ft.Colors.BLUE_200,
            focused_border_color=ft.Colors.BLUE_600,
            dense=True,
            expand=True
        )
        self.search_button = ft.ElevatedButton(
            "Tìm kiếm",
            icon=ft.Icons.SEARCH,
            on_click=self.filter_allowance_list,
            style=ft.ButtonStyle(
                bgcolor=ft.Colors.BLUE_800, 
                color=ft.Colors.WHITE, 
                shape=ft.RoundedRectangleBorder(radius=12),
                elevation=2
            )
        )
        self.list_container = ft.ListView(spacing=8, expand=True)

        # Cấu hình bộ lọc theo Ngày cho Tab 2
        self.target_date = datetime.now()
        self.date_picker = ft.DatePicker(
            on_change=self.change_target_date,
            first_date=datetime(1990, 1, 1),
            last_date=datetime(2100, 12, 31)
        )
        self.page.overlay.append(self.date_picker)

        # UI cho Tab 2: Bộ Lọc
        self.stats_search_field = ft.TextField(
            label="Tìm theo tên, số hiệu, đơn vị...",
            prefix_icon=ft.Icons.SEARCH,
            on_submit=self.apply_statistics_filter,
            border_radius=10,
            bgcolor=ft.Colors.WHITE,
            border_color=ft.Colors.BLUE_200,
            focused_border_color=ft.Colors.BLUE_600,
            dense=True,
            expand=True
        )
        
        self.btn_pick_date = ft.ElevatedButton(
            f"Tính đến: {self.target_date.strftime('%d/%m/%Y')}",
            icon=ft.Icons.CALENDAR_MONTH,
            on_click=lambda _: self.page.open(self.date_picker),
            style=ft.ButtonStyle(
                bgcolor=ft.Colors.WHITE,
                color=ft.Colors.BLUE_800,
                shape=ft.RoundedRectangleBorder(radius=10)
            )
        )

        self.filter_unit_dd = ft.Dropdown(label="Đơn vị", width=160, on_change=self.apply_statistics_filter, dense=True, border_radius=10, bgcolor=ft.Colors.WHITE)
        self.filter_rate_dd = ft.Dropdown(label="Mức hưởng", width=120, on_change=self.apply_statistics_filter, dense=True, border_radius=10, bgcolor=ft.Colors.WHITE)
        self.filter_status_dd = ft.Dropdown(
            label="Trạng thái",
            width=220,
            options=[
                ft.dropdown.Option("Tất cả"),
                ft.dropdown.Option("Đủ điều kiện chuyển mức"),
                ft.dropdown.Option("Còn <= 3 tháng chuyển mức"),
                ft.dropdown.Option("Đang hưởng an toàn"),
            ],
            on_change=self.apply_statistics_filter,
            dense=True,
            border_radius=10,
            bgcolor=ft.Colors.WHITE
        )
        self.export_docx_btn = ft.ElevatedButton(
            "Xuất Word",
            icon=ft.Icons.DESCRIPTION,
            on_click=self.trigger_docx_save_picker,
            style=ft.ButtonStyle(
                bgcolor=ft.Colors.BLUE_800,
                color=ft.Colors.WHITE,
                shape=ft.RoundedRectangleBorder(radius=10),
                elevation=2
            )
        )
        self.stats_total_text = ft.Text("Tổng số nhân sự: 0", size=14, weight=ft.FontWeight.BOLD, color=ft.Colors.BLUE_GREY_800)
        self.stats_container = ft.ListView(spacing=8, expand=True)

        # UI cho Tab 3: Thống Kê
        self.summary_stats_view = ft.ListView(spacing=10, expand=True)

        # UI cho Tab 5: Điều Chỉnh
        self.adjust_search_field = ft.TextField(
            label="Lọc nhanh nhân sự cần chỉnh sửa dữ liệu...",
            prefix_icon=ft.Icons.EDIT,
            on_change=self.filter_adjustment_list,
            border_radius=12,
            bgcolor=ft.Colors.WHITE,
            border_color=ft.Colors.TEAL_200,
            focused_border_color=ft.Colors.TEAL_600,
            dense=True,
            expand=True
        )
        self.adjust_table_container = ft.Column(scroll=ft.ScrollMode.ALWAYS, expand=True)

        self.create_ui()

        if os.path.exists("thuhut.xlsx"):
            self.process_excel_file("thuhut.xlsx")

    def change_target_date(self, e):
        if self.date_picker.value:
            self.target_date = self.date_picker.value
            self.btn_pick_date.text = f"Tính đến: {self.target_date.strftime('%d/%m/%Y')}"
            self.page.update()
            self.apply_statistics_filter(None)
            self.render_summary_stats()

    def create_ui(self):
        header = ft.Container(
            content=ft.Row([
                ft.Row([
                    ft.Icon(ft.Icons.LAYERS, color=ft.Colors.WHITE, size=28),
                    ft.Text("HỆ THỐNG QUẢN LÝ THU HÚT", size=18, weight=ft.FontWeight.BOLD, color=ft.Colors.WHITE),
                ], spacing=10),
                ft.ElevatedButton(
                    "Load Excel",
                    icon=ft.Icons.UPLOAD_FILE,
                    on_click=self.trigger_file_picker,
                    style=ft.ButtonStyle(
                        bgcolor=ft.Colors.AMBER_600, 
                        color=ft.Colors.WHITE, 
                        shape=ft.RoundedRectangleBorder(radius=10),
                        elevation=3
                    )
                )
            ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN),
            gradient=ft.LinearGradient(
                begin=ft.alignment.top_left,
                end=ft.alignment.bottom_right,
                colors=[ft.Colors.BLUE_900, ft.Colors.INDIGO_700]
            ),
            padding=ft.padding.symmetric(horizontal=20, vertical=15),
            border_radius=12,
            shadow=ft.BoxShadow(blur_radius=8, color=ft.Colors.BLUE_GREY_200, offset=ft.Offset(0, 3))
        )

        tab_danh_sach = ft.Container(
            content=ft.Column([
                ft.Row([self.search_field, self.search_button], spacing=10),
                self.build_table_header(),
                ft.Container(content=self.list_container, expand=True)
            ], spacing=10),
            padding=10
        )

        tab_bo_loc = ft.Container(
            content=ft.Column([
                ft.Container(
                    content=ft.Row([
                        self.stats_search_field, 
                        self.filter_unit_dd, 
                        self.filter_rate_dd, 
                        self.filter_status_dd,
                        self.btn_pick_date,
                        self.export_docx_btn
                    ], spacing=10, vertical_alignment=ft.CrossAxisAlignment.CENTER),
                    padding=10, bgcolor=ft.Colors.BLUE_GREY_100, border_radius=12
                ),
                ft.Row([
                    ft.Icon(ft.Icons.FILTER_ALT, color=ft.Colors.BLUE_700, size=20),
                    self.stats_total_text
                ], spacing=8),
                ft.Container(content=self.stats_container, expand=True)
            ], spacing=12),
            padding=10
        )

        tab_thong_ke = ft.Container(
            content=ft.Column([
                ft.Text("THỐNG KÊ CHI TIẾT THEO MỨC HƯỞNG VÀ MỐC TIẾP THEO", size=15, weight=ft.FontWeight.BOLD, color=ft.Colors.BLUE_900),
                ft.Text("💡 Số liệu được tự động tính toán dựa trên mốc 'Tính đến ngày' ở Tab 2.", size=12, italic=True, color=ft.Colors.TEAL_800),
                ft.Divider(height=1, color=ft.Colors.GREY_300),
                ft.Container(content=self.summary_stats_view, expand=True)
            ], spacing=15),
            padding=15
        )

        tab_huong_dan = ft.Container(
            content=ft.Column([
                ft.Text("VĂN BẢN HƯỚNG DẪN", size=15, weight=ft.FontWeight.BOLD, color=ft.Colors.BLUE_GREY_900),
                ft.Divider(height=1, color=ft.Colors.GREY_300),
            ], spacing=15, scroll=ft.ScrollMode.ALWAYS),
            padding=15
        )

        tab_dieu_chinh = ft.Container(
            content=ft.Column([
                ft.Row([self.adjust_search_field]),
                ft.Text("💡 Mẹo: Nhấp vào bất kỳ ô nào để trực tiếp sửa đổi.", size=12, italic=True, color=ft.Colors.TEAL_800),
                ft.Container(
                    content=self.adjust_table_container,
                    expand=True,
                    border=ft.border.all(1, ft.Colors.GREY_300),
                    border_radius=8,
                    bgcolor=ft.Colors.WHITE
                )
            ], spacing=10),
            padding=10
        )

        self.tabs_control = ft.Tabs(
            selected_index=0,
            animation_duration=200,
            tabs=[
                ft.Tab(text="1. Danh sách tổng hợp", icon=ft.Icons.LIST_ALT, content=tab_danh_sach),
                ft.Tab(text="2. Bộ lọc", icon=ft.Icons.FILTER_ALT, content=tab_bo_loc),
                ft.Tab(text="3. Thống kê", icon=ft.Icons.PIE_CHART, content=tab_thong_ke),
                ft.Tab(text="4. Văn bản hướng dẫn", icon=ft.Icons.BOOKMARK, content=tab_huong_dan),
                ft.Tab(text="5. Điều chỉnh", icon=ft.Icons.GRID_ON, content=tab_dieu_chinh),
            ],
            expand=True
        )

        self.page.add(ft.Column([header, ft.Container(height=5), self.tabs_control], expand=True, spacing=0))

    def get_cumulative_months(self, emp, target_dt):
        total_months = 0
        for rec in emp["all_records"]:
            s_str = rec.get("start_date", "")
            e_str = rec.get("end_date", "")
            if not s_str or s_str == "Chưa cập nhật": continue
            try:
                s_dt = datetime.strptime(s_str, '%d/%m/%Y')
                e_dt = target_dt
                if e_str and e_str != "Chưa cập nhật":
                    parsed_e = datetime.strptime(e_str, '%d/%m/%Y')
                    e_dt = min(target_dt, parsed_e)
                if e_dt > s_dt:
                    diff = relativedelta(e_dt, s_dt)
                    total_months += diff.years * 12 + diff.months
            except: pass
        return total_months

    def evaluate_cumulative_status(self, emp, target_dt):
        total_months = self.get_cumulative_months(emp, target_dt)
        if not emp["all_records"]:
            return "GREY", "Chưa rõ", ft.Colors.GREY_600
            
        latest_rec = emp["all_records"][-1]
        rate = str(latest_rec.get("rate", "70%")).strip().lower()
        
        # --- TRƯỜNG HỢP 1: ĐANG Ở MỨC THU HÚT 70% ---
        if "70" in rate:
            if total_months >= 120:
                # Đã vượt mốc 120 tháng
                return "RED", f"Đủ ĐK chuyển LN 0.7", ft.Colors.RED_600
            elif 60 <= total_months < 120:
                # Nằm trong khoảng từ 60 đến dưới 120 tháng -> Đang tiến về mốc 120 tháng
                left_to_120 = 120 - total_months
                if left_to_120 <= 3:
                    return "ORANGE", f"Còn {left_to_120} tháng chuyển LN 0.7", ft.Colors.ORANGE_700
                return "RED", f"Đủ ĐK chuyển LN 0.5", ft.Colors.RED_600
            else:
                # Dưới 60 tháng -> Đang tiến về mốc 60 tháng
                left_to_60 = 60 - total_months
                if left_to_60 <= 3:
                    return "ORANGE", f"Còn {left_to_60} tháng chuyển LN 0.5", ft.Colors.ORANGE_700
                return "GREEN", f"Còn {left_to_60} tháng chuyển LN 0.5", ft.Colors.GREEN_700
                
        # --- TRƯỜNG HỢP 2: ĐANG Ở MỨC LÂU NĂM 0.5 ---
        elif "0.5" in rate:
            if total_months >= 180:
                # Đã vượt mốc 180 tháng
                return "RED", f"Đủ ĐK chuyển LN 1.0", ft.Colors.RED_600
            elif 120 <= total_months < 180:
                # Nằm trong khoảng từ 120 đến dưới 180 tháng -> Đang tiến về mốc 180 tháng
                left_to_180 = 180 - total_months
                if left_to_180 <= 3:
                    return "ORANGE", f"Còn {left_to_180} tháng chuyển LN 1.0", ft.Colors.ORANGE_700
                return "RED", f"Đủ ĐK chuyển LN 0.7", ft.Colors.RED_600
            else:
                # Dưới 120 tháng -> Đang tiến về mốc 120 tháng
                left_to_120 = 120 - total_months
                if left_to_120 <= 3:
                    return "ORANGE", f"Đang LN 0.5 (còn {left_to_120} tháng lên 0.7)", ft.Colors.ORANGE_700
                return "GREEN", f"Đang LN 0.5 (còn {left_to_120} tháng lên 0.7)", ft.Colors.GREEN_700
                
        # --- TRƯỜNG HỢP 3: ĐANG Ở MỨC LÂU NĂM 0.7 ---
        elif "0.7" in rate:
            if total_months >= 180:
                # Đã vượt mốc 180 tháng
                return "RED", f"Đủ ĐK chuyển LN 1.0", ft.Colors.RED_600
            else:
                # Dưới 180 tháng -> Đang tiến về mốc 180 tháng
                left_to_180 = 180 - total_months
                if left_to_180 <= 3:
                    return "ORANGE", f"Đang LN 0.7 (còn {left_to_180} tháng lên 1.0)", ft.Colors.ORANGE_700
                return "GREEN", f"Đang LN 0.7 (còn {left_to_180} tháng lên 1.0)", ft.Colors.GREEN_700
                
        # --- TRƯỜNG HỢP 4: ĐANG Ở MỨC LÂU NĂM ĐỊCH TRẦN 1.0 ---
        elif "1.0" in rate or rate == "1" or "100" in rate:
            return "BLUE", f"Đang hưởng LN 1.0", ft.Colors.BLUE_700
            
        # --- TRƯỜNG HỢP 5: CHƯA RÕ MỨC HOẶC MỨC KHÁC ---
        else:
            if total_months >= 60:
                return "RED", f"Đủ ĐK chuyển LN 0.5", ft.Colors.RED_600
            else:
                left_to_60 = 60 - total_months
                if left_to_60 <= 3:
                    return "ORANGE", f"Còn {left_to_60} tháng chuyển LN 0.5", ft.Colors.ORANGE_700
                return "GREEN", f"Khác (còn {left_to_60} tháng lên 0.5)", ft.Colors.GREEN_700

    def evaluate_record_status(self, emp, rec_index, target_dt):
        # Nếu không phải là dòng ghi chú cuối cùng, báo là lịch sử đã qua
        if rec_index < len(emp["all_records"]) - 1:
            return "GREY", "Đã hoàn thành", ft.Colors.GREY_500
        return self.evaluate_cumulative_status(emp, target_dt)

    def render_summary_stats(self):
        self.summary_stats_view.controls.clear()

        # 1. Cấu trúc lưu trữ số lượng của đúng 9 mục yêu cầu
        stats = {
            "tong_doi_tuong": {"count": 0},
            "tong_hien_huong": {"count": 0},
            "dang_huong_70": {"count": 0},
            "dang_huong_05": {"count": 0},
            "dang_huong_07": {"count": 0},
            "dang_huong_10": {"count": 0},
            "du_dk_05": {"count": 0},
            "du_dk_07": {"count": 0},
            "du_dk_10": {"count": 0},
        }

        # 2. Quét dữ liệu master_data và phân loại số lượng khớp 100% với bộ lọc
        for norm_key, emp in self.master_data.items():
            if not emp["all_records"]: 
                continue
            
            latest_rec = emp["all_records"][-1]
            raw_rate = str(latest_rec.get("rate", "")).strip().lower().replace(",", ".")
            total_m = self.get_cumulative_months(emp, self.target_date)

            # - Mục 1: Tổng số đối tượng
            stats["tong_doi_tuong"]["count"] += 1

            # Kiểm tra trạng thái Hiện hưởng (chưa có ngày kết thúc)
            end_date_val = latest_rec.get("end_date")
            is_hien_huong = False
            
            if latest_rec.get("is_projected", False):
                is_hien_huong = True
            else:
                end_date_val = latest_rec.get("end_date")
                if not end_date_val:
                    is_hien_huong = True
                else:
                    end_date_str = str(end_date_val).strip().lower()
                    if end_date_str in ["", "nan", "nat", "-", "chưa cập nhật"]:
                        is_hien_huong = True

            # - Mục 2: Tổng số hiện hưởng
            if is_hien_huong:
                stats["tong_hien_huong"]["count"] += 1

            # - Mục 3 đến 6: Tổng số đang hưởng các mức
            if "70" in raw_rate:
                stats["dang_huong_70"]["count"] += 1
            elif "0.5" in raw_rate:
                stats["dang_huong_05"]["count"] += 1
            elif "0.7" in raw_rate:
                stats["dang_huong_07"]["count"] += 1
            elif "1.0" in raw_rate or raw_rate == "1" or "100" in raw_rate:
                stats["dang_huong_10"]["count"] += 1

            # - Mục 7 đến 9: Tính toán điều kiện đủ thâm niên chuyển mức hưởng
            if "70" in raw_rate:
                if total_m >= 120:
                    stats["du_dk_07"]["count"] += 1
                elif 60 <= total_m < 120:
                    stats["du_dk_05"]["count"] += 1
            elif "0.5" in raw_rate:
                if total_m >= 180:
                    stats["du_dk_10"]["count"] += 1
                elif 120 <= total_m < 180:
                    stats["du_dk_07"]["count"] += 1
            elif "0.7" in raw_rate:
                if total_m >= 180:
                    stats["du_dk_10"]["count"] += 1
            else:
                if total_m >= 60:
                    stats["du_dk_05"]["count"] += 1

        # 3. Định hình danh sách hiển thị phẳng theo cấu trúc 9 dòng yêu cầu
        ordered_config = [
            ("Tổng số đối tượng trong danh sách", "tong_doi_tuong", ft.Icons.PEOPLE_ALT),
            ("Tổng số hiện hưởng", "tong_hien_huong", ft.Icons.TOGGLE_ON),
            ("Tổng số đang hưởng mức 70%", "dang_huong_70", ft.Icons.PERCENT),
            ("Tổng số đang hưởng mức 0.5", "dang_huong_05", ft.Icons.LOOKS_5),
            ("Tổng số đang hưởng mức 0.7", "dang_huong_07", ft.Icons.LOOKS_6),
            ("Tổng số đang hưởng mức 1.0", "dang_huong_10", ft.Icons.LOOKS_ONE),
            ("Tổng số đủ điều kiện hưởng 0.5", "du_dk_05", ft.Icons.STAR_HALF),
            ("Tổng số đủ điều kiện hưởng 0.7", "du_dk_07", ft.Icons.STAR),
            ("Tổng số đủ điều kiện hưởng 1.0", "du_dk_10", ft.Icons.STAR),
        ]

        # 4. Tạo giao diện phẳng tĩnh (Static Row)
        for label, data_key, icon_name in ordered_config:
            count = stats[data_key]["count"]

            # Thiết lập màu sắc dựa trên việc có hay không có nhân sự
            theme_color = ft.Colors.BLUE_800 if count > 0 else ft.Colors.BLUE_GREY_300
            badge_bg = ft.Colors.BLUE_900 if count > 0 else ft.Colors.BLUE_GREY_400

            # Xây dựng container phẳng đơn giản hiển thị thông tin
            row_item = ft.Container(
                content=ft.Row([
                    ft.Icon(icon_name, color=theme_color, size=20),
                    ft.Text(label, size=13, weight=ft.FontWeight.W_500, expand=True),
                    ft.Container(
                        content=ft.Text(f"{count} người", weight=ft.FontWeight.BOLD, color=ft.Colors.WHITE, size=12),
                        bgcolor=badge_bg,
                        padding=ft.padding.symmetric(horizontal=12, vertical=4),
                        border_radius=12
                    )
                ], spacing=10),
                padding=12,
                bgcolor=ft.Colors.WHITE,
                border=ft.border.all(1, ft.Colors.GREY_200),
                border_radius=8,
                margin=ft.margin.only(bottom=6)
            )

            self.summary_stats_view.controls.append(row_item)

        self.page.update()

    def build_table_header(self):
        return ft.Container(
            content=ft.Row([
                ft.Text("Họ và tên", size=12, weight=ft.FontWeight.BOLD, width=180, color=ft.Colors.WHITE),
                ft.Text("Số hiệu", size=12, weight=ft.FontWeight.BOLD, width=100, color=ft.Colors.WHITE),
                ft.Text("Đơn vị hiện tại", size=12, weight=ft.FontWeight.BOLD, width=160, color=ft.Colors.WHITE),
                ft.Text("Mức hiện tại", size=12, weight=ft.FontWeight.BOLD, width=100, color=ft.Colors.WHITE),
                ft.Text("Tổng thời gian hưởng", size=12, weight=ft.FontWeight.BOLD, width=200, color=ft.Colors.WHITE), # <--- ĐÃ SỬA Ở ĐÂY
                ft.Text("Trạng thái", size=12, weight=ft.FontWeight.BOLD, expand=True, text_align=ft.TextAlign.RIGHT, color=ft.Colors.WHITE),
            ]),
            gradient=ft.LinearGradient(colors=[ft.Colors.BLUE_GREY_800, ft.Colors.BLUE_GREY_700]),
            padding=ft.padding.symmetric(horizontal=12, vertical=12),
            border_radius=8
        )

    def trigger_file_picker(self, e):
        file_picker = ft.FilePicker(on_result=self.process_excel_file)
        self.page.overlay.append(file_picker)
        self.page.update()
        file_picker.pick_files(allowed_extensions=["xlsx"])

    def process_excel_file(self, e):
        if e is None: return
        try:
            # Nếu e là một chuỗi đường dẫn trực tiếp (tự động load)
            if isinstance(e, str):
                path = e
            else:
                # Ngược lại nếu gọi qua nút bấm FilePicker
                if not e.files: return
                path = e.files[0].path
            
            sheets_dict = pd.read_excel(path, sheet_name=None)
            self.master_data.clear()
            self.extra_headers.clear()

            for sheet_name, df in sheets_dict.items():
                if df.empty: continue
                df.columns = [str(c).strip() for c in df.columns]

                name_col = find_column_by_keywords(df.columns, ['ho va ten', 'ho ten', 'hoten', 'ten'])
                yob_col = find_column_by_keywords(df.columns, ['nam sinh', 'namsinh', 'ngay sinh', 'ngaysinh'])
                cand_col = find_column_by_keywords(df.columns, ['so hieu cand', 'so hieu', 'sh cand', 'so hieu quan nhan'])
                unit_col = find_column_by_keywords(df.columns, ['don vi hien tai', 'don vi', 'phong ban', 'donvi'])
                rate_col = find_column_by_keywords(df.columns, ['muc huong', 'ty le', 'he so', 'muc phu cap'])
                start_date_col = find_column_by_keywords(df.columns, ['bat dau huong', 'thoi gian nhan cong tactu ngay', 'tu ngay'])
                end_date_col = find_column_by_keywords(df.columns, ['ket thuc huong', 'han huong', 'den ngay'])
                total_months_col = find_column_by_keywords(df.columns, ['tong so thang', 'tong thang', 'so thang', 'thoi gian'])

                if not name_col: continue
                last_name, last_yob, last_cand, last_unit = "", "Chưa cập nhật", "Chưa cập nhật", "Chưa cập nhật"

                yob_regex = re.compile(r"\((19\d{2}|20\d{2})\)")
                yob_replace_regex = re.compile(r"\s*\(19\d{2}\)|\s*\(20\d{2}\)")
                today_date = datetime.now() 

                for row_tuple in df.itertuples(index=False):
                    row = dict(zip(df.columns, row_tuple))

                    raw_name = str(row[name_col]).strip() if pd.notna(row[name_col]) else ""
                    
                    extracted_yob = ""
                    if raw_name and raw_name.lower() != 'nan':
                        match_yob = yob_regex.search(raw_name) 
                        if match_yob:
                            extracted_yob = match_yob.group(1)
                            raw_name = yob_replace_regex.sub("", raw_name).strip()

                    raw_yob = ""
                    if yob_col and pd.notna(row[yob_col]):
                        val_yob = str(row[yob_col]).strip()
                        if val_yob.endswith('.0'): val_yob = val_yob[:-2]
                        raw_yob = val_yob
                    
                    if not raw_yob or raw_yob.lower() in ['nan', '', 'nat', '-', '.', '0']:
                        raw_yob = extracted_yob if extracted_yob else "Chưa cập nhật"

                    if raw_name and raw_name.lower() != 'nan':
                        last_name = raw_name
                        last_yob = raw_yob if (raw_yob and raw_yob.lower() != 'nan') else "Chưa cập nhật"
                        last_cand = "Chưa cập nhật"
                        last_unit = "Chưa cập nhật"
                    else:
                        raw_name = last_name
                        raw_yob = last_yob

                    if not raw_name or raw_name.lower() in ['nan', '']: continue

                    norm_key = f"{normalize_text(raw_name)}_{normalize_text(raw_yob)}"

                    raw_cand = str(row[cand_col]).strip() if cand_col and pd.notna(row[cand_col]) else ""
                    if raw_cand and raw_cand.lower() != 'nan': last_cand = raw_cand
                    else: raw_cand = last_cand

                    raw_unit = str(row[unit_col]).strip() if unit_col and pd.notna(row[unit_col]) else ""
                    if raw_unit and raw_unit.lower() != 'nan': last_unit = clean_unit_name(raw_unit)
                    else: raw_unit = last_unit

                    # Lấy Mức hưởng thô
                    rate_val = str(row[rate_col]).strip() if rate_col and pd.notna(row[rate_col]) else ""

                    start_val = clean_cell_to_date_str(row[start_date_col]) if start_date_col else ""
                    end_val = clean_cell_to_date_str(row[end_date_col]) if end_date_col else ""

                    # Khôi phục logic 5 năm và thêm biến is_projected
                    is_projected = False 
                    if start_val and not end_val:
                        try:
                            start_dt = datetime.strptime(start_val, '%d/%m/%Y')
                            end_dt = start_dt + relativedelta(years=5)
                            end_val = end_dt.strftime('%d/%m/%Y')
                            is_projected = True # Đánh dấu đây là ngày dự kiến
                        except: pass

                    smart_months_val = "Chưa cập nhật"
                    if start_val and start_val != "Chưa cập nhật":
                        try:
                            start_dt = datetime.strptime(start_val, '%d/%m/%Y')
                            end_point = today_date
                            if end_val and end_val != "Chưa cập nhật":
                                end_dt = datetime.strptime(end_val, '%d/%m/%Y')
                                end_point = min(today_date, end_dt)

                            if end_point >= start_dt:
                                diff = relativedelta(end_point, start_dt)
                                total_m = diff.years * 12 + diff.months
                                breakdown = []
                                if diff.years > 0: breakdown.append(f"{diff.years} năm")
                                if diff.months > 0: breakdown.append(f"{diff.months} tháng")
                                smart_months_val = f"{total_m} tháng ({' '.join(breakdown)})" if breakdown else f"{total_m} tháng"
                        except: pass

                    row_details = {}
                    unnamed_idx = 1
                    chuthich_lines = [] 

                    for k, v in row.items():
                        k_str = str(k).strip()
                        is_unnamed = "unnamed" in k_str.lower() or not k_str
                        v_str = str(v).strip()
                        has_value = not (pd.isna(v) or v_str.lower() in ['nan', 'nat', '-', '.', '0', ''])

                        if is_unnamed:
                            if has_value: chuthich_lines.append(f"- Ghi chú {unnamed_idx}: {v_str}")
                            unnamed_idx += 1
                            continue

                        if name_col and k == name_col: row_details[k_str] = raw_name
                        elif yob_col and k == yob_col: row_details[k_str] = raw_yob
                        elif cand_col and k == cand_col: row_details[k_str] = raw_cand
                        elif unit_col and k == unit_col: row_details[k_str] = raw_unit
                        elif rate_col == k: row_details[k_str] = rate_val
                        elif start_date_col and k == start_date_col: row_details[k_str] = start_val if start_val else "Chưa cập nhật"
                        elif end_date_col and k == end_date_col: row_details[k_str] = end_val if end_val else "Chưa cập nhật"
                        elif total_months_col and k == total_months_col: row_details[k_str] = smart_months_val
                        else:
                            row_details[k_str] = v_str if has_value else "Chưa cập nhật"

                        core_cols = [
                            str(name_col).strip() if name_col else "",
                            str(yob_col).strip() if yob_col else "",
                            str(cand_col).strip() if cand_col else "",
                            str(unit_col).strip() if unit_col else "",
                            str(rate_col).strip() if rate_col else "",
                            str(start_date_col).strip() if start_date_col else "",
                            str(end_date_col).strip() if end_date_col else ""
                        ]
                        if k_str not in core_cols and k_str not in self.extra_headers:
                            self.extra_headers.append(k_str)

                    if chuthich_lines: row_details["Chú thích"] = "\n".join(chuthich_lines)
                    else: row_details["Chú thích"] = "Không có"

                    if "Chú thích" not in self.extra_headers: self.extra_headers.append("Chú thích")

                    if not total_months_col:
                        row_details["Tổng số tháng"] = smart_months_val
                        if "Tổng số tháng" not in self.extra_headers:
                            self.extra_headers.append("Tổng số tháng")

                    if start_val and end_val: duration_text = f"{start_val} → {end_val}"
                    elif start_val: duration_text = f"Từ {start_val}"
                    elif end_val: duration_text = f"Đến {end_val}"
                    else: duration_text = "Chưa cập nhật"

                    if norm_key not in self.master_data:
                        self.master_data[norm_key] = {
                            "display_name": raw_name,
                            "yob": raw_yob,
                            "cand_id": raw_cand if raw_cand else "Chưa cập nhật",
                            "unit": raw_unit if raw_unit else "Chưa cập nhật",
                            "all_records": []
                        }

                    is_duplicate = False
                    for existing in self.master_data[norm_key]["all_records"]:
                        if (existing["start_date"] == start_val and 
                            existing["end_date"] == end_val and 
                            existing["rate"] == rate_val and 
                            existing["unit"] == (raw_unit if raw_unit else "Chưa cập nhật")):
                            is_duplicate = True
                            break
                    
                    if not is_duplicate:
                        self.master_data[norm_key]["all_records"].append({
                            "sheet": sheet_name,
                            "start_date": start_val,
                            "end_date": end_val,
                            "is_projected": is_projected, # THÊM DÒNG NÀY
                            "rate": rate_val,
                            "duration_str": duration_text,
                            "full_row": row_details,
                            "unit": raw_unit if raw_unit else "Chưa cập nhật",
                            "cand_id": raw_cand if raw_cand else "Chưa cập nhật",
                        })

            # Gom các record của người dùng bị phân tách
            name_to_keys = {}
            for k, emp in self.master_data.items():
                norm_name = normalize_text(emp["display_name"])
                if norm_name not in name_to_keys:
                    name_to_keys[norm_name] = []
                name_to_keys[norm_name].append(k)

            for norm_name, keys in name_to_keys.items():
                if len(keys) > 1:
                    with_yob, without_yob = [], []
                    for k in keys:
                        yob_clean = str(self.master_data[k]["yob"]).strip().lower()
                        if yob_clean in ["chưa cập nhật", "nan", ""]: without_yob.append(k)
                        else: with_yob.append(k)

                    if len(with_yob) == 1 and len(without_yob) > 0:
                        target_key = with_yob[0]
                        for w_k in without_yob:
                            self.master_data[target_key]["all_records"].extend(self.master_data[w_k]["all_records"])
                            del self.master_data[w_k]
                    elif len(with_yob) == 0 and len(without_yob) > 1:
                        target_key = without_yob[0]
                        for w_k in without_yob[1:]:
                            self.master_data[target_key]["all_records"].extend(self.master_data[w_k]["all_records"])
                            del self.master_data[w_k]

            def get_sort_date(rec):
                s_date = rec.get("start_date", "")
                if not s_date or s_date == "Chưa cập nhật": return datetime.min
                try: return datetime.strptime(s_date, "%d/%m/%Y")
                except: return datetime.min

            # Sắp xếp lịch sử quá trình và Cập nhật Mức hưởng thừa kế (Inheritance)
            for norm_key, emp in self.master_data.items():
                emp["all_records"].sort(key=get_sort_date)
                
                final_unit, final_cand = "Chưa cập nhật", "Chưa cập nhật"
                last_rate = "70%" # Mặc định dòng đầu
                
                for rec in emp["all_records"]:
                    u = rec.get("unit", "Chưa cập nhật")
                    if u and u != "Chưa cập nhật" and final_unit == "Chưa cập nhật": final_unit = u
                    c = rec.get("cand_id", "Chưa cập nhật")
                    if c and c != "Chưa cập nhật" and final_cand == "Chưa cập nhật": final_cand = c
                    
                    # Logic thừa kế và chuẩn hóa mức hưởng
                    raw_rate = str(rec["full_row"].get("Mức hiện tại") or rec.get("rate", "")).strip().lower()
                    if not raw_rate or raw_rate in ['nan', '', '-', '.', '0', 'chưa cập nhật']:
                        n_rate = last_rate
                    elif "70" in raw_rate: n_rate = "70%"
                    elif "0.5" in raw_rate: n_rate = "0.5"
                    elif "0.7" in raw_rate: n_rate = "0.7"
                    elif "1.0" in raw_rate or raw_rate == "1" or "100" in raw_rate: n_rate = "1.0"
                    else: n_rate = str(rec.get("rate", "")) if str(rec.get("rate", "")) else last_rate
                    
                    rec["rate"] = n_rate
                    rec["full_row"]["Mức hiện tại"] = n_rate
                    last_rate = n_rate
                    
                emp["unit"] = final_unit
                emp["cand_id"] = final_cand

            unique_units = set(emp["unit"] for emp in self.master_data.values() if emp["unit"] != "Chưa cập nhật")
            unique_rates = set(rec["rate"] for emp in self.master_data.values() for rec in emp["all_records"] if rec["rate"] != "Chưa cập nhật")

            self.filter_unit_dd.options = [ft.dropdown.Option("Tất cả")] + [
                ft.dropdown.Option(u) for u in sorted(unique_units, key=lambda x: str(x).lower())
            ]
            self.filter_rate_dd.options = [ft.dropdown.Option("Tất cả")] + [
                ft.dropdown.Option(r) for r in sorted(unique_rates, key=lambda x: str(x).lower())
            ]
            self.filter_unit_dd.value = "Tất cả"
            self.filter_rate_dd.value = "Tất cả"
            self.filter_status_dd.value = "Tất cả"

            self.filtered_data = self.master_data.copy()
            self.render_allowance_list(append=False)
            self.render_statistics_list(append=False)
            self.render_summary_stats()
            self.render_adjustment_list(append=False)

            self.page.open(ft.SnackBar(ft.Text("Đồng bộ dữ liệu thành công!"), bgcolor=ft.Colors.GREEN_600))
        except Exception as ex:
            self.page.open(ft.SnackBar(ft.Text(f"Lỗi đọc file: {str(ex)}"), bgcolor=ft.Colors.RED_600))
        self.page.update()

    def get_clean_display_title(self, emp) -> str:
        name = emp.get("display_name", "")
        yob = str(emp.get("yob", "")).strip()
        if yob and yob.lower() not in ["chưa cập nhật", "nan", ""]: 
            return f"{name} ({yob})"
        return name

    def show_history_popup(self, person_key):
        emp = self.master_data.get(person_key)
        if not emp or not emp["all_records"]: return

        display_yob = emp['yob'] if emp['yob'] != "Chưa cập nhật" else ""
        
        # Tính toán tổng thời gian hiển thị ở Popup
        total_m = self.get_cumulative_months(emp, datetime.now())
        years = total_m // 12
        months = total_m % 12
        breakdown = []
        if years > 0: breakdown.append(f"{years} năm")
        if months > 0: breakdown.append(f"{months} tháng")
        total_duration_str = f"{total_m} tháng ({' '.join(breakdown)})" if breakdown else "0 tháng"

        # Bổ sung cột "Tổng TG hưởng" bên cạnh "Đơn vị hiện tại"
        info_profile = ft.Container(
            content=ft.ResponsiveRow([
                ft.Column([
                    ft.Text("Họ và tên", size=11, color=ft.Colors.BLUE_GREY_400, weight=ft.FontWeight.W_500),
                    ft.Text(emp['display_name'], size=14, weight=ft.FontWeight.BOLD, color=ft.Colors.BLUE_900)
                ], col={"xs": 12, "sm": 6, "md": 3}),
                ft.Column([
                    ft.Text("Năm sinh", size=11, color=ft.Colors.BLUE_GREY_400, weight=ft.FontWeight.W_500),
                    ft.Text(display_yob, size=14, weight=ft.FontWeight.BOLD, color=ft.Colors.BLUE_900)
                ], col={"xs": 12, "sm": 6, "md": 2}),
                ft.Column([
                    ft.Text("Số hiệu CAND", size=11, color=ft.Colors.BLUE_GREY_400, weight=ft.FontWeight.W_500),
                    ft.Text(emp['cand_id'], size=14, weight=ft.FontWeight.BOLD, color=ft.Colors.BLUE_900)
                ], col={"xs": 12, "sm": 6, "md": 2}),
                ft.Column([
                    ft.Text("Đơn vị hiện tại", size=11, color=ft.Colors.BLUE_GREY_400, weight=ft.FontWeight.W_500),
                    ft.Text(emp['unit'], size=14, weight=ft.FontWeight.BOLD, color=ft.Colors.BLUE_900, overflow=ft.TextOverflow.ELLIPSIS)
                ], col={"xs": 12, "sm": 6, "md": 3}),
                ft.Column([
                    ft.Text("Tổng thời gian hưởng", size=11, color=ft.Colors.BLUE_GREY_400, weight=ft.FontWeight.W_500),
                    ft.Text(total_duration_str, size=14, weight=ft.FontWeight.BOLD, color=ft.Colors.RED_600)
                ], col={"xs": 12, "sm": 6, "md": 2}),
            ], spacing=10),
            bgcolor=ft.Colors.BLUE_GREY_50, padding=15, border_radius=8, margin=ft.margin.only(bottom=15)
        )

        sample_row = emp["all_records"][0]["full_row"]
        all_headers = list(sample_row.keys())
        exclude_kws = ['ho va ten', 'ho ten', 'hoten', 'ten', 'nam sinh', 'namsinh', 'ngay sinh', 'ngaysinh', 'so hieu cand', 'so hieu', 'sh cand', 'so hieu quan nhan']
        filtered_headers = [h for h in all_headers if normalize_text(h) not in exclude_kws]

        extended_headers = filtered_headers + ["Trạng thái"]
        data_columns = [ft.DataColumn(ft.Text(h, weight=ft.FontWeight.BOLD, color=ft.Colors.BLUE_GREY_700, size=12)) for h in extended_headers]

        data_rows = []
        for idx, rec in enumerate(emp["all_records"]):
            # Lấy trạng thái của bản ghi dựa trên thời điểm hiện tại
            _, badge_desc, color_flet = self.evaluate_record_status(emp, idx, datetime.now())

            row_cells = []
            for h in filtered_headers:
                v = rec["full_row"].get(h, "Chưa cập nhật")
                row_cells.append(ft.DataCell(ft.Text(str(v), size=12, color=ft.Colors.BLUE_GREY_900)))

            row_cells.append(ft.DataCell(
                ft.Container(
                    content=ft.Text(badge_desc, color=ft.Colors.WHITE, size=10, weight=ft.FontWeight.BOLD),
                    bgcolor=color_flet, padding=ft.padding.symmetric(horizontal=8, vertical=3), border_radius=4
                )
            ))
            data_rows.append(ft.DataRow(cells=row_cells))

        detail_table = ft.DataTable(
            columns=data_columns, rows=data_rows, heading_row_color=ft.Colors.GREY_100, divider_thickness=1,
            horizontal_lines=ft.BorderSide(0.5, ft.Colors.GREY_300), vertical_lines=ft.BorderSide(0.5, ft.Colors.GREY_200),
        )

        scrollable_container = ft.Column(expand=True, scroll=ft.ScrollMode.ALWAYS, controls=[ft.Row(scroll=ft.ScrollMode.ALWAYS, controls=[detail_table])])

        popup_layout = ft.Column([
            info_profile,
            ft.Text("📅 Lịch sử quá trình:", size=12, weight=ft.FontWeight.BOLD, color=ft.Colors.BLUE_GREY_600),
            ft.Container(content=scrollable_container, expand=True)
        ], spacing=5, expand=True)

        def redirect_to_excel_mode(e):
            self.page.close(dialog)
            self.tabs_control.selected_index = 4 
            self.adjust_search_field.value = emp['display_name']
            self.render_adjustment_list(append=False)
            self.page.update()

        dialog = ft.AlertDialog(
            title=ft.Row([
                ft.Icon(ft.Icons.HISTORY, color=ft.Colors.BLUE_800, size=24),
                ft.Text("Chi tiết", size=16, weight=ft.FontWeight.BOLD)
            ], spacing=8),
            content=ft.Container(content=popup_layout, width=1050, height=520),
            actions=[
                ft.ElevatedButton("Điều chỉnh/Cập nhật thông tin", on_click=redirect_to_excel_mode, style=ft.ButtonStyle(bgcolor=ft.Colors.TEAL_700, color=ft.Colors.WHITE, shape=ft.RoundedRectangleBorder(radius=8))),
                ft.TextButton("Đóng ", on_click=lambda e: self.page.close(dialog), style=ft.ButtonStyle(color=ft.Colors.BLUE_GREY_700))
            ]
        )
        self.page.open(dialog)

    def load_more_data(self, e):
        self.page_number += 1
        self.render_allowance_list(append=True)

    def render_allowance_list(self, append=False):
        if not append:
            self.list_container.controls.clear()
            self.page_number = 1

        if not self.filtered_data:
            self.list_container.controls.append(ft.Text("Không tìm thấy kết quả phù hợp.", color=ft.Colors.GREY_500, italic=True))
            self.page_number = 1
            self.page.update()
            return

        if append and len(self.list_container.controls) > 0:
            if getattr(self.list_container.controls[-1], "key", None) == "load_more_btn":
                self.list_container.controls.pop()

        items = list(self.filtered_data.items())
        start_idx = (self.page_number - 1) * self.per_page
        end_idx = start_idx + self.per_page
        page_items = items[start_idx:end_idx]

        for norm_key, emp in page_items:
            # 1. Tính tổng thời gian đã hưởng
            total_m = self.get_cumulative_months(emp, datetime.now())
            years = total_m // 12
            months = total_m % 12
            breakdown = []
            if years > 0: breakdown.append(f"{years} năm")
            if months > 0: breakdown.append(f"{months} tháng")
            duration_str = f"{total_m} tháng ({' '.join(breakdown)})" if breakdown else "0 tháng"
            
            # Giao diện hiển thị tổng thời gian
            duration_view = ft.Text(duration_str, size=13, weight=ft.FontWeight.W_500, color=ft.Colors.TEAL_800)

            latest_rec = emp["all_records"][-1]
            
            # 2. Tính toán tình trạng Lũy Kế dựa theo thời điểm hiện tại (Ngày chạy App)
            c_code, badge_desc, color_flet = self.evaluate_cumulative_status(emp, datetime.now())
            display_title = self.get_clean_display_title(emp)

            row_card = ft.Container(
                content=ft.Row([
                    ft.Container(content=ft.Text(display_title, size=13, weight=ft.FontWeight.BOLD, color=ft.Colors.BLUE_800, overflow=ft.TextOverflow.ELLIPSIS), width=180),
                    ft.Text(emp["cand_id"], size=13, width=100),
                    ft.Text(emp["unit"], size=13, width=160, overflow=ft.TextOverflow.ELLIPSIS),
                    ft.Container(content=ft.Text(latest_rec["rate"], size=13), width=100),
                    
                    # 3. Thay thế ft.Column cũ bằng Container chứa tổng thời gian
                    ft.Container(content=duration_view, width=200),
                    
                    ft.Container(
                        content=ft.Container(
                            content=ft.Text(badge_desc, color=ft.Colors.WHITE, size=11, weight=ft.FontWeight.W_500),
                            bgcolor=color_flet, padding=ft.padding.symmetric(horizontal=10, vertical=2), border_radius=4,
                        ),
                        expand=True, alignment=ft.alignment.center_right
                    )
                ], vertical_alignment=ft.CrossAxisAlignment.CENTER),
                bgcolor=ft.Colors.WHITE, padding=ft.padding.symmetric(horizontal=12, vertical=12), border_radius=8,
                border=ft.border.all(0.5, ft.Colors.GREY_200), on_click=lambda e, k=norm_key: self.show_history_popup(k),
                shadow=ft.BoxShadow(blur_radius=4, color=ft.Colors.GREY_200, offset=ft.Offset(0, 1))
            )
            self.list_container.controls.append(row_card)

        if len(items) > end_idx:
            load_more_btn = ft.Container(
                key="load_more_btn",
                content=ft.ElevatedButton(
                    f"Xem thêm ({len(items) - end_idx} nhân sự trong kết quả tìm kiếm)...",
                    icon=ft.Icons.ARROW_DROP_DOWN_CIRCLE_ROUNDED,
                    on_click=self.load_more_data,
                    style=ft.ButtonStyle(bgcolor=ft.Colors.BLUE_50, color=ft.Colors.BLUE_800)
                ),
                alignment=ft.alignment.center, padding=ft.padding.only(top=10, bottom=10)
            )
            self.list_container.controls.append(load_more_btn)
        self.page.update()

    def filter_allowance_list(self, e):
        q = normalize_text(self.search_field.value)
        if not q: self.filtered_data = self.master_data.copy()
        else:
            self.filtered_data = {}
            for k, emp in self.master_data.items():
                if (q in normalize_text(emp["display_name"]) or q in normalize_text(emp["cand_id"]) or q in normalize_text(emp["unit"]) or q in normalize_text(emp["yob"])):
                    self.filtered_data[k] = emp
        self.render_allowance_list(append=False)

    def load_more_stats_data(self, e):
        self.stats_page_number += 1
        self.render_statistics_list(append=True)

    def render_statistics_list(self, append=False):
        if not append:
            self.stats_container.controls.clear()
            self.stats_page_number = 1
            self.stats_filtered_items = []

            q_stats = normalize_text(self.stats_search_field.value)
            unit_f = self.filter_unit_dd.value
            rate_f = self.filter_rate_dd.value
            status_f = self.filter_status_dd.value

            for norm_key, emp in self.master_data.items():
                if not emp["all_records"]: continue
                
                if q_stats:
                    if not (q_stats in normalize_text(emp["display_name"]) or 
                            q_stats in normalize_text(emp["cand_id"]) or 
                            q_stats in normalize_text(emp["unit"]) or 
                            q_stats in normalize_text(emp["yob"])):
                        continue

                if unit_f and unit_f != "Tất cả" and emp["unit"] != unit_f: continue
                
                match_found = True
                if (rate_f and rate_f != "Tất cả") or (status_f and status_f != "Tất cả"):
                    has_rate, has_status = False, False
                    
                    latest_rec = emp["all_records"][-1]
                    if rate_f and rate_f != "Tất cả" and latest_rec["rate"] == rate_f: has_rate = True
                    
                    color_code, _, _ = self.evaluate_cumulative_status(emp, self.target_date)
                    if status_f and status_f != "Tất cả":
                        if status_f == "Đủ điều kiện chuyển mức" and color_code == "RED": has_status = True
                        if status_f == "Còn <= 3 tháng chuyển mức" and color_code == "ORANGE": has_status = True
                        if status_f == "Đang hưởng an toàn" and color_code in ["GREEN", "BLUE"]: has_status = True

                    if rate_f and rate_f != "Tất cả" and not has_rate: match_found = False
                    if status_f and status_f != "Tất cả" and not has_status: match_found = False

                if not match_found: continue
                self.stats_filtered_items.append((norm_key, emp))
            self.stats_total_text.value = f"📊 Tổng số nhân sự thỏa mãn điều kiện: {len(self.stats_filtered_items)} người"

        if append and len(self.stats_container.controls) > 0:
            if getattr(self.stats_container.controls[-1], "key", None) == "load_more_stats_btn":
                self.stats_container.controls.pop()

        start_idx = (self.stats_page_number - 1) * self.per_page
        end_idx = start_idx + self.per_page
        page_items = self.stats_filtered_items[start_idx:end_idx]

        for norm_key, emp in page_items:
            latest_rec = emp["all_records"][-1]
            
            # Badge tự động tính theo `target_date` tại Tab Thống kê
            _, badge_desc, color_flet = self.evaluate_cumulative_status(emp, self.target_date)
            display_title = self.get_clean_display_title(emp)

            stats_row = ft.Container(
                content=ft.Row([
                    ft.Text(display_title, size=13, weight=ft.FontWeight.W_500, width=180),
                    ft.Text(emp["unit"], size=13, width=200),
                    ft.Container(content=ft.Text(latest_rec["rate"], size=13), width=100),
                    ft.Container(
                        content=ft.Container(
                            content=ft.Text(badge_desc, color=ft.Colors.WHITE, size=10),
                            bgcolor=color_flet, padding=ft.padding.symmetric(horizontal=8, vertical=2), border_radius=4
                        ),
                        expand=True, alignment=ft.alignment.center_right
                    )
                ], vertical_alignment=ft.CrossAxisAlignment.CENTER),
                bgcolor=ft.Colors.WHITE, padding=12, border_radius=6,
                on_click=lambda e, k=norm_key: self.show_history_popup(k),
                border=ft.border.all(0.5, ft.Colors.GREY_200)
            )
            self.stats_container.controls.append(stats_row)

        if len(self.stats_filtered_items) > end_idx:
            load_more_stats_btn = ft.Container(
                key="load_more_stats_btn",
                content=ft.ElevatedButton(
                    f"Xem thêm ({len(self.stats_filtered_items) - end_idx} nhân sự)...",
                    icon=ft.Icons.ARROW_DROP_DOWN_CIRCLE_ROUNDED,
                    on_click=self.load_more_stats_data,
                    style=ft.ButtonStyle(bgcolor=ft.Colors.BLUE_50, color=ft.Colors.BLUE_800)
                ),
                alignment=ft.alignment.center, padding=ft.padding.only(top=10, bottom=10)
            )
            self.stats_container.controls.append(load_more_stats_btn)
        self.page.update()

    def apply_statistics_filter(self, e):
        self.render_statistics_list(append=False)

    def trigger_docx_save_picker(self, e):
        if not self.stats_filtered_items:
            self.page.open(ft.SnackBar(ft.Text("Không có dữ liệu thỏa mãn bộ lọc hiện tại để xuất Word!"), bgcolor=ft.Colors.ORANGE_600))
            return
        
        docx_picker = ft.FilePicker(on_result=self.generate_docx_callback)
        self.page.overlay.append(docx_picker)
        self.page.update()
        docx_picker.save_file(
            allowed_extensions=["docx"], 
            file_name=f"Bao_cao_phu_cap_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        )

    def generate_docx_callback(self, e: ft.FilePickerResultEvent):
        if not e.path: return
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_ALIGN_VERTICAL

            doc = Document()
            
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.6)
                section.bottom_margin = Inches(0.6)
                section.left_margin = Inches(0.6)
                section.right_margin = Inches(0.6)

            num_people = len(self.stats_filtered_items)

            if num_people == 1:
                _, emp = self.stats_filtered_items[0]
                
                header_table = doc.add_table(rows=1, cols=2)
                header_table.autofit = False
                header_table.columns[0].width = Inches(3.0)
                header_table.columns[1].width = Inches(4.5)
                
                cell_left = header_table.cell(0, 0)
                p_left = cell_left.paragraphs[0]
                p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_l1 = p_left.add_run("CÔNG AN TỈNH CAO BẰNG\n")
                r_l1.font.name = 'Arial'
                r_l1.font.size = Pt(11)
                r_l2 = p_left.add_run("PHÒNG TỔ CHỨC CÁN BỘ")
                r_l2.font.name = 'Arial'
                r_l2.font.size = Pt(11)
                r_l2.font.bold = True
                r_l2.font.underline = True
                
                cell_right = header_table.cell(0, 1)
                p_right = cell_right.paragraphs[0]
                p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_r1 = p_right.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n")
                r_r1.font.name = 'Arial'
                r_r1.font.size = Pt(11)
                r_r1.font.bold = True
                r_r2 = p_right.add_run("Độc lập – Tự do – Hạnh phúc")
                r_r2.font.name = 'Arial'
                r_r2.font.size = Pt(11)
                r_r2.font.bold = True
                r_r2.font.underline = True
                
                p_space = doc.add_paragraph()
                p_space.paragraph_format.space_before = Pt(12)
                
                p_date = doc.add_paragraph()
                p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                current_date_str = f"Cao Bằng, ngày {datetime.now().day} tháng {datetime.now().month} năm {datetime.now().year}"
                r_date = p_date.add_run(current_date_str)
                r_date.font.name = 'Arial'
                r_date.font.size = Pt(11)
                r_date.font.italic = True
                
                p_space2 = doc.add_paragraph()
                p_space2.paragraph_format.space_before = Pt(18)
                
                p_title = doc.add_paragraph()
                p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_title = p_title.add_run("QUÁ TRÌNH HƯỞNG PHỤ CẤP")
                r_title.font.name = 'Arial'
                r_title.font.size = Pt(14)
                r_title.font.bold = True
                
                p_space3 = doc.add_paragraph()
                p_space3.paragraph_format.space_before = Pt(18)
                
                sh_str = emp['cand_id'] if emp['cand_id'] != "Chưa cập nhật" else ""
                p_info1 = doc.add_paragraph()
                r_info1 = p_info1.add_run(f"Họ và tên: {emp['display_name']};        Số hiệu: {sh_str}")
                r_info1.font.name = 'Arial'
                r_info1.font.size = Pt(11)
                
                dv_str = emp['unit'] if emp['unit'] != "Chưa cập nhật" else ""
                p_info2 = doc.add_paragraph()
                r_info2 = p_info2.add_run(f"Đơn vị: {dv_str}")
                r_info2.font.name = 'Arial'
                r_info2.font.size = Pt(11)
                
                p_info3 = doc.add_paragraph()
                r_info3 = p_info3.add_run("Quá trình hưởng:")
                r_info3.font.name = 'Arial'
                r_info3.font.size = Pt(11)
                r_info3.font.bold = True
                
                num_records = len(emp["all_records"])
                table = doc.add_table(rows=num_records + 2, cols=6, style='Table Grid')
                table.autofit = False
                
                col_widths = [Inches(0.5), Inches(2.2), Inches(1.1), Inches(1.1), Inches(1.5), Inches(1.1)]
                for row in table.rows:
                    for idx, width in enumerate(col_widths):
                        row.cells[idx].width = width
                
                headers = ["STT", "Đơn vị", "Thời gian bắt đầu", "Thời gian kết thúc", "Lý do kết thúc", "Thời gian"]
                for idx, h_text in enumerate(headers):
                    cell = table.cell(0, idx)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run(h_text)
                    r.font.name = 'Arial'
                    r.font.size = Pt(10)
                    r.font.bold = True
                
                total_months_sum = 0
                for idx, rec in enumerate(emp["all_records"]):
                    row_idx = idx + 1
                    stt = str(idx + 1)
                    unit = rec.get("unit", "")
                    if unit == "Chưa cập nhật": unit = ""
                    start = rec.get("start_date", "")
                    if start == "Chưa cập nhật": start = ""
                    end = rec.get("end_date", "")
                    if end == "Chưa cập nhật": end = ""
                    
                    reason = rec["full_row"].get("Lý do kết thúc", rec["full_row"].get("Lý do", ""))
                    if reason == "Chưa cập nhật": reason = ""
                        
                    duration = rec["full_row"].get("Tổng số tháng", rec["full_row"].get("Thời gian", ""))
                    if duration == "Chưa cập nhật": duration = ""
                    
                    match = re.search(r"(\d+)\s*tháng", str(duration).lower())
                    if match:
                        total_months_sum += int(match.group(1))
                    elif str(duration).isdigit():
                        total_months_sum += int(duration)
                        duration = f"{duration} tháng"
                    
                    vals = [stt, str(unit), str(start), str(end), str(reason), str(duration)]
                    alignments = [
                        WD_ALIGN_PARAGRAPH.CENTER,
                        WD_ALIGN_PARAGRAPH.LEFT,
                        WD_ALIGN_PARAGRAPH.CENTER,
                        WD_ALIGN_PARAGRAPH.CENTER,
                        WD_ALIGN_PARAGRAPH.LEFT,
                        WD_ALIGN_PARAGRAPH.CENTER
                    ]
                    
                    for col_idx, val in enumerate(vals):
                        cell = table.cell(row_idx, col_idx)
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        p = cell.paragraphs[0]
                        p.alignment = alignments[col_idx]
                        r = p.add_run(val)
                        r.font.name = 'Arial'
                        r.font.size = Pt(10)
                
                footer_row_idx = num_records + 1
                total_duration_str = ""
                if total_months_sum > 0:
                    years = total_months_sum // 12
                    months = total_months_sum % 12
                    breakdown = []
                    if years > 0: breakdown.append(f"{years} năm")
                    if months > 0: breakdown.append(f"{months} tháng")
                    total_duration_str = f"{total_months_sum} tháng" + (f" ({' '.join(breakdown)})" if breakdown else "")
                
                cell_start = table.cell(footer_row_idx, 0)
                cell_end = table.cell(footer_row_idx, 4)
                merged_cell = cell_start.merge(cell_end)
                merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p_foot1 = merged_cell.paragraphs[0]
                p_foot1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_foot1 = p_foot1.add_run("Tổng thời gian")
                r_foot1.font.name = 'Arial'
                r_foot1.font.size = Pt(10)
                r_foot1.font.bold = True
                
                cell_last = table.cell(footer_row_idx, 5)
                cell_last.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p_foot2 = cell_last.paragraphs[0]
                p_foot2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_foot2 = p_foot2.add_run(total_duration_str)
                r_foot2.font.name = 'Arial'
                r_foot2.font.size = Pt(10)
                r_foot2.font.bold = True

            else:
                header_table = doc.add_table(rows=1, cols=2)
                header_table.autofit = False
                header_table.columns[0].width = Inches(3.0)
                header_table.columns[1].width = Inches(4.5)
                
                cell_left = header_table.cell(0, 0)
                p_left = cell_left.paragraphs[0]
                p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_l1 = p_left.add_run("CÔNG AN TỈNH CAO BẰNG\n")
                r_l1.font.name = 'Arial'
                r_l1.font.size = Pt(11)
                r_l2 = p_left.add_run("PHÒNG TỔ CHỨC CÁN BỘ")
                r_l2.font.name = 'Arial'
                r_l2.font.size = Pt(11)
                r_l2.font.bold = True
                r_l2.font.underline = True
                
                cell_right = header_table.cell(0, 1)
                p_right = cell_right.paragraphs[0]
                p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_r1 = p_right.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n")
                r_r1.font.name = 'Arial'
                r_r1.font.size = Pt(11)
                r_r1.font.bold = True
                r_r2 = p_right.add_run("Độc lập – Tự do – Hạnh phúc")
                r_r2.font.name = 'Arial'
                r_r2.font.size = Pt(11)
                r_r2.font.bold = True
                r_r2.font.underline = True
                
                p_space = doc.add_paragraph()
                p_space.paragraph_format.space_before = Pt(12)
                
                p_date = doc.add_paragraph()
                p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                current_date_str = f"Cao Bằng, ngày {datetime.now().day} tháng {datetime.now().month} năm {datetime.now().year}"
                r_date = p_date.add_run(current_date_str)
                r_date.font.name = 'Arial'
                r_date.font.size = Pt(11)
                r_date.font.italic = True
                
                p_space2 = doc.add_paragraph()
                p_space2.paragraph_format.space_before = Pt(18)
                
                p_title = doc.add_paragraph()
                p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r_title = p_title.add_run("DANH SÁCH TỔNG HỢP QUÁ TRÌNH HƯỞNG PHỤ CẤP")
                r_title.font.name = 'Arial'
                r_title.font.size = Pt(14)
                r_title.font.bold = True
                
                p_space3 = doc.add_paragraph()
                p_space3.paragraph_format.space_before = Pt(18)
                
                total_rows_needed = sum(len(emp["all_records"]) for _, emp in self.stats_filtered_items)
                
                table = doc.add_table(rows=total_rows_needed + 1, cols=8, style='Table Grid')
                table.autofit = False
                
                col_widths = [Inches(0.4), Inches(1.4), Inches(0.8), Inches(1.3), Inches(0.9), Inches(0.9), Inches(0.9), Inches(0.8)]
                for row in table.rows:
                    for idx, width in enumerate(col_widths):
                        row.cells[idx].width = width
                
                headers = ["STT", "Họ và tên", "Số hiệu", "Đơn vị công tác", "TG Bắt đầu", "TG Kết thúc", "Lý do kết thúc", "Thời gian"]
                for idx, h_text in enumerate(headers):
                    cell = table.cell(0, idx)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run(h_text)
                    r.font.name = 'Arial'
                    r.font.size = Pt(10)
                    r.font.bold = True
                
                global_stt = 1
                row_idx = 1
                for _, emp in self.stats_filtered_items:
                    sh_str = emp['cand_id'] if emp['cand_id'] != "Chưa cập nhật" else ""
                    num_records = len(emp["all_records"])
                    start_row = row_idx  
                    
                    for i, rec in enumerate(emp["all_records"]):
                        unit = rec.get("unit", "")
                        if unit == "Chưa cập nhật": unit = ""
                        start = rec.get("start_date", "")
                        if start == "Chưa cập nhật": start = ""
                        end = rec.get("end_date", "")
                        if end == "Chưa cập nhật": end = ""
                        
                        reason = rec["full_row"].get("Lý do kết thúc", rec["full_row"].get("Lý do", ""))
                        if reason == "Chưa cập nhật": reason = ""
                            
                        duration = rec["full_row"].get("Tổng số tháng", rec["full_row"].get("Thời gian", ""))
                        if duration == "Chưa cập nhật": duration = ""
                        if str(duration).isdigit():
                            duration = f"{duration} tháng"
                        
                        vals = [str(global_stt), emp['display_name'], sh_str, str(unit), str(start), str(end), str(reason), str(duration)]
                        alignments = [
                            WD_ALIGN_PARAGRAPH.CENTER,
                            WD_ALIGN_PARAGRAPH.LEFT,
                            WD_ALIGN_PARAGRAPH.CENTER,
                            WD_ALIGN_PARAGRAPH.LEFT,
                            WD_ALIGN_PARAGRAPH.CENTER,
                            WD_ALIGN_PARAGRAPH.CENTER,
                            WD_ALIGN_PARAGRAPH.LEFT,
                            WD_ALIGN_PARAGRAPH.CENTER
                        ]
                        
                        for col_idx, val in enumerate(vals):
                            if i > 0 and col_idx in [0, 1, 2]:
                                continue
                                
                            cell = table.cell(row_idx, col_idx)
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            p = cell.paragraphs[0]
                            p.alignment = alignments[col_idx]
                            r = p.add_run(val)
                            r.font.name = 'Arial'
                            r.font.size = Pt(9.5)
                        
                        row_idx += 1
                    
                    if num_records > 1:
                        end_row = start_row + num_records - 1
                        for col_idx in [0, 1, 2]: 
                            cell_start = table.cell(start_row, col_idx)
                            cell_end = table.cell(end_row, col_idx)
                            cell_start.merge(cell_end)
                            
                    global_stt += 1 

            doc.save(e.path)
            self.page.open(ft.SnackBar(ft.Text(f"Xuất file biểu mẫu Word thành công tại: {e.path}"), bgcolor=ft.Colors.GREEN_600))
        except Exception as ex:
            self.page.open(ft.SnackBar(ft.Text(f"Lỗi hệ thống xuất Word: {str(ex)}"), bgcolor=ft.Colors.RED_600))
        self.page.update()

    def filter_adjustment_list(self, e):
        self.render_adjustment_list(append=False)

    def load_more_adjust_data(self, e):
        self.adjust_page_number += 1
        self.render_adjustment_list(append=True)

    def render_adjustment_list(self, append=False):
        if not append:
            self.adjust_table_container.controls.clear()
            self.adjust_page_number = 1
            
            self.adjust_grid = ft.ListView(spacing=0, expand=True) 
            
            self.adjust_horizontal_wrapper = ft.Row(
                scroll=ft.ScrollMode.ALWAYS, 
                controls=[
                    ft.Container(content=self.adjust_grid, width=1500, expand=True) 
                ],
                expand=True
            )
            self.adjust_table_container.controls.append(self.adjust_horizontal_wrapper)
        else:
            if len(self.adjust_grid.controls) > 0:
                if getattr(self.adjust_grid.controls[-1], "key", None) == "more_adjust_btn_container":
                    self.adjust_grid.controls.pop()
        
        q = normalize_text(self.adjust_search_field.value)
        
        matching_items = []
        for k, emp in self.master_data.items():
            if not q or (q in normalize_text(emp["display_name"]) or q in normalize_text(emp["cand_id"]) or q in normalize_text(emp["unit"])):
                matching_items.append((k, emp))

        if not matching_items:
            self.adjust_grid.controls.append(
                ft.Padding(padding=20, content=ft.Text("Không tìm thấy dữ liệu nhân sự phù hợp để sửa.", color=ft.Colors.GREY_500, italic=True))
            )
            self.page.update()
            return

        max_name_len = max([len(emp["display_name"]) for _, emp in matching_items] + [10])
        max_cand_len = max([len(str(emp["cand_id"])) for _, emp in matching_items] + [8])
        max_unit_len = 10
        for _, emp in matching_items:
            for rec in emp["all_records"]:
                max_unit_len = max(max_unit_len, len(str(rec["unit"])))

        w_name = min(180, max(100, int(max_name_len * 6.5) + 15))
        w_yob = 70      
        w_cand = max(80, int(max_cand_len * 6.5) + 15)
        w_unit = min(250, max(120, int(max_unit_len * 6.5) + 15))
        w_rate = 75     
        w_start = 85    
        w_end = 85      
        w_extra = 130 

        row_height = 38 
        header_height = 35

        grid_rows = []
        
        if not append:
            header_controls = [
                ft.Container(content=ft.Text("Họ tên", weight=ft.FontWeight.BOLD, color=ft.Colors.TEAL_900, size=12), width=w_name, height=header_height, bgcolor=ft.Colors.TEAL_50, alignment=ft.alignment.center, border=ft.border.all(0.5, ft.Colors.GREY_300)),
                ft.Container(content=ft.Text("Năm sinh", weight=ft.FontWeight.BOLD, color=ft.Colors.TEAL_900, size=12), width=w_yob, height=header_height, bgcolor=ft.Colors.TEAL_50, alignment=ft.alignment.center, border=ft.border.all(0.5, ft.Colors.GREY_300)),
                ft.Container(content=ft.Text("Số hiệu", weight=ft.FontWeight.BOLD, color=ft.Colors.TEAL_900, size=12), width=w_cand, height=header_height, bgcolor=ft.Colors.TEAL_50, alignment=ft.alignment.center, border=ft.border.all(0.5, ft.Colors.GREY_300)),
                ft.Container(content=ft.Text("Đơn vị", weight=ft.FontWeight.BOLD, color=ft.Colors.TEAL_900, size=12), width=w_unit, height=header_height, bgcolor=ft.Colors.TEAL_50, alignment=ft.alignment.center, border=ft.border.all(0.5, ft.Colors.GREY_300)),
                ft.Container(content=ft.Text("Mức", weight=ft.FontWeight.BOLD, color=ft.Colors.TEAL_900, size=12), width=w_rate, height=header_height, bgcolor=ft.Colors.TEAL_50, alignment=ft.alignment.center, border=ft.border.all(0.5, ft.Colors.GREY_300)),
                ft.Container(content=ft.Text("Bắt đầu", weight=ft.FontWeight.BOLD, color=ft.Colors.TEAL_900, size=12), width=w_start, height=header_height, bgcolor=ft.Colors.TEAL_50, alignment=ft.alignment.center, border=ft.border.all(0.5, ft.Colors.GREY_300)),
                ft.Container(content=ft.Text("Kết thúc", weight=ft.FontWeight.BOLD, color=ft.Colors.TEAL_900, size=12), width=w_end, height=header_height, bgcolor=ft.Colors.TEAL_50, alignment=ft.alignment.center, border=ft.border.all(0.5, ft.Colors.GREY_300)),
            ]
            
            for ext_header in self.extra_headers:
                header_controls.append(
                    ft.Container(content=ft.Text(ext_header, weight=ft.FontWeight.BOLD, color=ft.Colors.TEAL_900, size=12), width=w_extra, height=header_height, bgcolor=ft.Colors.TEAL_50, alignment=ft.alignment.center, border=ft.border.all(0.5, ft.Colors.GREY_300))
                )

            grid_rows.append(ft.Row(header_controls, spacing=0))

        start_idx = (self.adjust_page_number - 1) * self.per_page if append else 0
        end_idx = self.adjust_page_number * self.per_page
        page_items = matching_items[start_idx:end_idx]

        for emp_key, emp in page_items:
            num_records = len(emp["all_records"])
            total_emp_height = num_records * row_height

            def make_cell_changer(e_key, r_idx, field_type, extra_key=None):
                def cell_on_change(e):
                    val = e.control.value
                    target_emp = self.master_data[e_key]
                    target_rec = target_emp["all_records"][r_idx]
                    if field_type == "name": target_emp["display_name"] = val
                    elif field_type == "yob": target_emp["yob"] = val if val.strip() else "Chưa cập nhật"
                    elif field_type == "cand_id": target_emp["cand_id"] = val
                    elif field_type == "unit": 
                        target_rec["unit"] = val
                        target_emp["unit"] = val 
                    elif field_type == "rate": target_rec["rate"] = val
                    elif field_type == "start": target_rec["start_date"] = val
                    elif field_type == "end": target_rec["end_date"] = val
                    elif field_type == "extra" and extra_key:
                        target_rec["full_row"][extra_key] = val
                return cell_on_change

            yob_display = emp["yob"] if emp["yob"] != "Chưa cập nhật" else ""
            cell_tf_padding = ft.padding.symmetric(horizontal=4, vertical=2)

            name_tf = ft.TextField(value=emp["display_name"], border=ft.InputBorder.NONE, dense=True, text_size=12, text_align=ft.TextAlign.CENTER, content_padding=cell_tf_padding, on_change=make_cell_changer(emp_key, 0, "name"))
            yob_tf = ft.TextField(value=yob_display, border=ft.InputBorder.NONE, dense=True, text_size=12, text_align=ft.TextAlign.CENTER, content_padding=cell_tf_padding, on_change=make_cell_changer(emp_key, 0, "yob"))
            cand_tf = ft.TextField(value=emp["cand_id"], border=ft.InputBorder.NONE, dense=True, text_size=12, text_align=ft.TextAlign.CENTER, content_padding=cell_tf_padding, on_change=make_cell_changer(emp_key, 0, "cand_id"))

            record_subrows = []
            for idx, rec in enumerate(emp["all_records"]):
                unit_tf = ft.TextField(value=rec["unit"], border=ft.InputBorder.NONE, dense=True, text_size=12, content_padding=cell_tf_padding, on_change=make_cell_changer(emp_key, idx, "unit"))
                rate_tf = ft.TextField(value=rec["rate"], border=ft.InputBorder.NONE, dense=True, text_size=12, text_align=ft.TextAlign.CENTER, content_padding=cell_tf_padding, on_change=make_cell_changer(emp_key, idx, "rate"))
                start_tf = ft.TextField(value=rec["start_date"], border=ft.InputBorder.NONE, dense=True, text_size=12, text_align=ft.TextAlign.CENTER, content_padding=cell_tf_padding, on_change=make_cell_changer(emp_key, idx, "start"))
                end_tf = ft.TextField(value=rec["end_date"], border=ft.InputBorder.NONE, dense=True, text_size=12, text_align=ft.TextAlign.CENTER, content_padding=cell_tf_padding, on_change=make_cell_changer(emp_key, idx, "end"))

                subrow_controls = [
                    ft.Container(content=unit_tf, width=w_unit, height=row_height, alignment=ft.alignment.center_left, border=ft.border.only(bottom=ft.BorderSide(0.5, ft.Colors.GREY_300), right=ft.BorderSide(0.5, ft.Colors.GREY_300))),
                    ft.Container(content=rate_tf, width=w_rate, height=row_height, alignment=ft.alignment.center, border=ft.border.only(bottom=ft.BorderSide(0.5, ft.Colors.GREY_300), right=ft.BorderSide(0.5, ft.Colors.GREY_300))),
                    ft.Container(content=start_tf, width=w_start, height=row_height, alignment=ft.alignment.center, border=ft.border.only(bottom=ft.BorderSide(0.5, ft.Colors.GREY_300), right=ft.BorderSide(0.5, ft.Colors.GREY_300))),
                    ft.Container(content=end_tf, width=w_end, height=row_height, alignment=ft.alignment.center, border=ft.border.only(bottom=ft.BorderSide(0.5, ft.Colors.GREY_300), right=ft.BorderSide(0.5, ft.Colors.GREY_300)))
                ]

                for ext_header in self.extra_headers:
                    ext_val = rec["full_row"].get(ext_header, "Chưa cập nhật")
                    ext_tf = ft.TextField(value=ext_val, border=ft.InputBorder.NONE, dense=True, text_size=12, text_align=ft.TextAlign.LEFT, content_padding=cell_tf_padding, on_change=make_cell_changer(emp_key, idx, "extra", ext_header))
                    subrow_controls.append(
                        ft.Container(content=ext_tf, width=w_extra, height=row_height, alignment=ft.alignment.center_left, border=ft.border.only(bottom=ft.BorderSide(0.5, ft.Colors.GREY_300), right=ft.BorderSide(0.5, ft.Colors.GREY_300)))
                    )

                record_subrows.append(ft.Row(subrow_controls, spacing=0))

            emp_block_row = ft.Row([
                ft.Container(content=name_tf, width=w_name, height=total_emp_height, alignment=ft.alignment.center, border=ft.border.only(bottom=ft.BorderSide(0.5, ft.Colors.GREY_300), right=ft.BorderSide(0.5, ft.Colors.GREY_300))),
                ft.Container(content=yob_tf, width=w_yob, height=total_emp_height, alignment=ft.alignment.center, border=ft.border.only(bottom=ft.BorderSide(0.5, ft.Colors.GREY_300), right=ft.BorderSide(0.5, ft.Colors.GREY_300))),
                ft.Container(content=cand_tf, width=w_cand, height=total_emp_height, alignment=ft.alignment.center, border=ft.border.only(bottom=ft.BorderSide(0.5, ft.Colors.GREY_300), right=ft.BorderSide(0.5, ft.Colors.GREY_300))),
                ft.Column(controls=record_subrows, spacing=0)
            ], spacing=0)

            grid_rows.append(emp_block_row)

        self.adjust_grid.controls.extend(grid_rows)

        if len(matching_items) > end_idx:
            more_btn = ft.Container(
                key="more_adjust_btn_container",
                content=ft.ElevatedButton(
                    f"Tải thêm ({len(matching_items) - end_idx} dòng tiếp theo)...",
                    icon=ft.Icons.GRID_ON,
                    on_click=self.load_more_adjust_data,
                    style=ft.ButtonStyle(bgcolor=ft.Colors.TEAL_50, color=ft.Colors.TEAL_900)
                ),
                alignment=ft.alignment.center, padding=ft.padding.all(10)
            )
            self.adjust_grid.controls.append(more_btn)

        self.page.update()

def main(page: ft.Page):
    AllowanceDashboardApp(page)

if __name__ == "__main__":
    ft.app(target=main)
